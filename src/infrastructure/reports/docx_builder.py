"""Сборка DOCX-файла аналитического отчёта из готовых данных."""
import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def build_report_docx(data: dict) -> BytesIO:
    """Построить DOCX-документ из JSON-данных аналитического отчёта.

    Args:
        data: словарь, соответствующий AnalyticalReportResponseSchema.

    Returns:
        BytesIO с готовым .docx файлом.
    """
    doc = Document()

    # Базовый стиль шрифта
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # Заголовок документа
    title = doc.add_heading("Аналитический отчёт по охране труда", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    today = datetime.date.today().strftime("%d.%m.%Y")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Дата формирования: {today}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 1. Сводная статистика
    if data.get("summary_narrative"):
        doc.add_heading("Сводная статистика", level=2)
        for para_text in data["summary_narrative"].split("\n"):
            stripped = para_text.strip()
            if stripped:
                doc.add_paragraph(stripped)

    # 2. Ключевые выводы
    findings = data.get("key_findings") or []
    if findings:
        doc.add_heading("Ключевые выводы", level=2)
        for item in findings:
            doc.add_paragraph(item, style="List Bullet")

    # 3. Анализ причин
    if data.get("cause_analysis"):
        doc.add_heading("Анализ причин", level=2)
        for para_text in data["cause_analysis"].split("\n"):
            stripped = para_text.strip()
            if stripped:
                doc.add_paragraph(stripped)

    # 4. Категории причин
    categories = data.get("top_cause_categories") or []
    if categories:
        doc.add_heading("Категории причин", level=2)
        for cat in categories:
            p = doc.add_paragraph(style="List Bullet")
            name = cat.get("category", "")
            count = cat.get("count", 0)
            analysis = cat.get("analysis", "")
            run_b = p.add_run(f"{name} ({count})")
            run_b.bold = True
            if analysis:
                p.add_run(f" — {analysis}")

    # 5. Паттерны повторяемости
    patterns = data.get("recurrence_patterns") or []
    if patterns:
        doc.add_heading("Паттерны повторяемости", level=2)
        for pat in patterns:
            p = doc.add_paragraph(style="List Bullet")
            desc = pat.get("pattern_description", "")
            companies = pat.get("affected_companies") or []
            text = desc
            if companies:
                text += f" ({', '.join(companies)})"
            p.add_run(text)

    # 6. Оценка рисков
    risks = data.get("risk_assessment") or []
    level = data.get("overall_risk_level")
    if risks or level:
        doc.add_heading("Оценка рисков", level=2)
        if level:
            p = doc.add_paragraph()
            p.add_run("Общий уровень: ").bold = True
            p.add_run(level)
        for risk in risks:
            p = doc.add_paragraph(style="List Bullet")
            rtype = risk.get("risk_type", "")
            severity = risk.get("severity", "")
            description = risk.get("description", "")
            run_b = p.add_run(f"{rtype} [{severity}]")
            run_b.bold = True
            p.add_run(f" — {description}")

    # 7. Рекомендации
    recs = data.get("recommendations") or []
    if recs:
        doc.add_heading("Рекомендации", level=2)
        for rec in recs:
            p = doc.add_paragraph(style="List Bullet")
            priority = rec.get("priority", "")
            text = rec.get("recommendation", "")
            rationale = rec.get("rationale", "")
            run_b = p.add_run(f"[{priority}] ")
            run_b.bold = True
            p.add_run(text)
            if rationale:
                p.add_run("\n")
                run_i = p.add_run(rationale)
                run_i.italic = True

    # 8. Первоочередные действия
    actions = data.get("immediate_actions") or []
    if actions:
        doc.add_heading("Первоочередные действия", level=2)
        for action in actions:
            doc.add_paragraph(action, style="List Number")

    # 9. Примечание
    if data.get("confidence_note"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        run_i = p.add_run(data["confidence_note"])
        run_i.italic = True
        run_i.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
